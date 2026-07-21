from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/urieeldev/Documents/Codes/portfolio-urieeldev/CV-Uriel-Mojica-Mejia-Actualizado.docx")

PURPLE = "4F24E8"
DARK_PURPLE = "271766"
INK = "171524"
MUTED = "625F70"
LIGHT_PURPLE = "EEE9FF"
LIGHT_RULE = "D8D1F5"
WHITE = "FFFFFF"
FONT = "Arial"
CONTENT_WIDTH_DXA = 10224  # 7.1 in, from 0.7 in left/right margins.


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_dxa, indent_dxa=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=INK, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before=0, after=4, line=1.12, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    return paragraph


def add_hyperlink(paragraph, text, url, size=9.2, color=PURPLE, bold=False):
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_props.append(fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    run_props.append(color_el)
    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size * 2)))
    run_props.append(size_el)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(size * 2)))
    run_props.append(size_cs)
    if bold:
        run_props.append(OxmlElement("w:b"))
    run.append(run_props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_bottom_border(paragraph, color=LIGHT_RULE, size=8, space=3):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "420")
    indent.set(qn("w:hanging"), "260")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "50")
    spacing.set(qn("w:line"), "260")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PURPLE)
    r_pr.append(color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def assign_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_body(doc, text, size=10, color=INK, bold=False, italic=False, after=4, keep=False):
    paragraph = style_paragraph(doc.add_paragraph(), after=after, line=1.12, keep_with_next=keep)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_section_heading(doc, text, before=10):
    paragraph = style_paragraph(doc.add_paragraph(), before=before, after=6, line=1.0, keep_with_next=True)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=12, color=PURPLE, bold=True)
    add_bottom_border(paragraph)
    return paragraph


def add_position_header(doc, company, period, role, location):
    paragraph = style_paragraph(doc.add_paragraph(), after=1, line=1.0, keep_with_next=True)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), alignment=WD_TAB_ALIGNMENT.RIGHT)
    company_run = paragraph.add_run(company)
    set_run_font(company_run, size=11, color=DARK_PURPLE, bold=True)
    paragraph.add_run("\t")
    period_run = paragraph.add_run(period)
    set_run_font(period_run, size=9.2, color=MUTED, bold=True)

    sub = style_paragraph(doc.add_paragraph(), after=4, line=1.0, keep_with_next=True)
    role_run = sub.add_run(role)
    set_run_font(role_run, size=9.7, color=INK, bold=True)
    location_run = sub.add_run(f"  ·  {location}")
    set_run_font(location_run, size=9.2, color=MUTED)


def add_bullet(doc, text, num_id, size=9.8):
    paragraph = style_paragraph(doc.add_paragraph(), after=4, line=1.12)
    assign_bullet(paragraph, num_id)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=INK)
    return paragraph


def add_project(doc, title, category, url, display_url, description):
    header = style_paragraph(doc.add_paragraph(), before=4, after=1, line=1.0, keep_with_next=True)
    title_run = header.add_run(title)
    set_run_font(title_run, size=10.5, color=DARK_PURPLE, bold=True)
    category_run = header.add_run(f"  ·  {category}  ·  ")
    set_run_font(category_run, size=8.8, color=MUTED, bold=True)
    add_hyperlink(header, display_url, url, size=8.8, color=PURPLE, bold=True)
    body = add_body(doc, description, size=9.7, color=INK, after=6)
    body.paragraph_format.keep_together = True


def set_footer(section):
    footer = section.footer
    footer.distance = Inches(0.3)
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run("urieel.dev  ·  Uriel Mojica Mejía  ·  ")
    set_run_font(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    r_pr.append(sz)
    r.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    r.append(text)
    fld.append(r)
    p._p.append(fld)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.58)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.3)
set_footer(section)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.12

for style_name, size, color, before, after in (
    ("Title", 28, DARK_PURPLE, 0, 1),
    ("Subtitle", 12.5, MUTED, 0, 3),
    ("Heading 1", 12, PURPLE, 10, 6),
    ("Heading 2", 10.5, DARK_PURPLE, 4, 1),
    ("Heading 3", 9.7, INK, 2, 1),
):
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = style_name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.keep_with_next = True

bullet_num_id = create_bullet_numbering(doc)

# Header metadata grid: deliberate two-column contact block.
header_table = doc.add_table(rows=1, cols=2)
set_table_geometry(header_table, [6200, 4024], indent_dxa=0)
remove_table_borders(header_table)
header_table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
header_table.rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
for cell in header_table.rows[0].cells:
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

left = header_table.rows[0].cells[0]
p = style_paragraph(left.paragraphs[0], after=1, line=1.0)
r = p.add_run("Uriel Mojica Mejía")
set_run_font(r, size=28, color=PURPLE, bold=True)
p = style_paragraph(left.add_paragraph(), after=3, line=1.0)
r = p.add_run("Senior Software Engineer")
set_run_font(r, size=12.5, color=INK, bold=True)
p = style_paragraph(left.add_paragraph(), after=0, line=1.0)
r = p.add_run("Frontend-first  ·  Full-stack  ·  E-commerce")
set_run_font(r, size=9.4, color=MUTED, bold=True)

right = header_table.rows[0].cells[1]
contact_lines = [
    ("México", None),
    ("+52 55 3113 9560", "tel:+525531139560"),
    ("urieel.dev@gmail.com", "mailto:urieel.dev@gmail.com"),
    ("urieel.dev", "https://urieel.dev"),
    ("linkedin.com/in/uriel-mm", "https://www.linkedin.com/in/uriel-mm"),
]
for index, (label, url) in enumerate(contact_lines):
    p = right.paragraphs[0] if index == 0 else right.add_paragraph()
    style_paragraph(p, after=1.2, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if url:
        add_hyperlink(p, label, url, size=8.9, color=PURPLE, bold=index == 1)
    else:
        r = p.add_run(label)
        set_run_font(r, size=8.9, color=MUTED, bold=True)

tagline = add_body(
    doc,
    "Senior Software Engineer con más de 7 años desarrollando productos digitales y plataformas de e-commerce de alto rendimiento. Especialización profunda en frontend y experiencia full-stack en servicios backend, datos, automatizaciones e integraciones cloud.",
    size=10.4,
    color=INK,
    after=4,
)
tagline.paragraph_format.space_before = Pt(7)

p = style_paragraph(doc.add_paragraph(), after=5, line=1.0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
for index, item in enumerate(("7+ años de experiencia", "+20% en ventas", "+50% en tráfico", "Frontend + Backend + Cloud")):
    if index:
        sep = p.add_run("   ·   ")
        set_run_font(sep, size=9.1, color=LIGHT_RULE, bold=True)
    run = p.add_run(item)
    set_run_font(run, size=9.3, color=DARK_PURPLE, bold=True)

add_section_heading(doc, "Experiencia profesional", before=6)
add_position_header(doc, "Corebiz México", "Noviembre 2020 – Julio 2026", "Software Engineer", "México")
add_body(
    doc,
    "Desarrollo, mantenimiento y optimización de plataformas de comercio electrónico para marcas nacionales e internacionales, en colaboración con equipos de diseño, producto, backend, marketing y negocio.",
    size=9.8,
    color=MUTED,
    after=4,
)

professional_bullets = [
    "Desarrollé y optimicé soluciones para Walmart Centroamérica, Miniso, HEB, La Colonia, Vianney, Óptima, Devlyn, Mi Tienda del Ahorro, Chanel, Avante, Supermercados La Torre, Maxidespensa y Plaforama.",
    "Construí componentes y funcionalidades reutilizables con React.js, Next.js, TypeScript, JavaScript, Sass y VTEX IO, orientados a escalabilidad, rendimiento y experiencia de usuario.",
    "Implementé componentes en React para Walmart Connect e integraciones con campañas y anuncios dinámicos de Google Ads.",
    "Rediseñé y optimicé el checkout de HEB, mejorando la interfaz, la navegación y la experiencia durante el proceso de compra.",
    "Desarrollé componentes de alta complejidad para Miniso México, Colombia, Chile y Perú, adaptables a los requerimientos comerciales y técnicos de cada país.",
    "Implementé mejoras de UX y automatización que contribuyeron a incrementar hasta en un 20 % las ventas; las integraciones con Google Tag Manager y Meta Pixel contribuyeron a aumentar hasta en un 50 % el tráfico.",
    "Desarrollé un módulo de facturación, middlewares y servicios con Node.js para conectar aplicaciones VTEX IO con sistemas internos y plataformas externas.",
]
for item in professional_bullets:
    add_bullet(doc, item, bullet_num_id)

p = style_paragraph(doc.add_paragraph(), before=3, after=0, line=1.0)
r = p.add_run("Stack principal: ")
set_run_font(r, size=9.2, color=DARK_PURPLE, bold=True)
r = p.add_run("React.js · Next.js · TypeScript · JavaScript · Node.js · VTEX IO · Sass")
set_run_font(r, size=9.2, color=MUTED)

# Page 2
page_break = doc.add_paragraph()
page_break.add_run().add_break(WD_BREAK.PAGE)

add_section_heading(doc, "Proyectos independientes seleccionados", before=0)
add_body(
    doc,
    "Productos desarrollados con alcances definidos, asumiendo responsabilidades end-to-end de frontend, backend, datos e integraciones.",
    size=9.8,
    color=MUTED,
    after=4,
)

add_project(
    doc,
    "EstateAdmin",
    "Plataforma SaaS full-stack",
    "https://estate-admin.com/",
    "estate-admin.com",
    "Desarrollé de forma full-stack un sistema integral para la gestión financiera, operativa y de comunicación de condominios. Incluye pagos, gastos, presupuestos, morosidad, reportes, tickets de mantenimiento, proyectos, reservaciones, roles, exportaciones PDF/Excel y notificaciones automáticas mediante WhatsApp y correo electrónico.",
)
add_project(
    doc,
    "YW Studio",
    "Programa de lealtad",
    "https://admin.ywstudio.com.mx/",
    "admin.ywstudio.com.mx",
    "Sistema web de programa de lealtad para un estudio de baile y cafetería, desarrollado como producto independiente para centralizar la operación del programa y la experiencia de sus usuarios.",
)
add_project(
    doc,
    "CEHF",
    "Educación a distancia",
    "https://cehf.live/",
    "cehf.live",
    "Plataforma digital para ofrecer experiencias de educación a distancia y facilitar el acceso a contenidos de aprendizaje en línea.",
)

add_section_heading(doc, "Capacidades técnicas", before=7)
skill_rows = [
    ("Frontend & UX", "React.js, Next.js, TypeScript, JavaScript, Tailwind CSS, Sass, sistemas de componentes, accesibilidad y performance."),
    ("Backend & APIs", "Node.js, NestJS, REST APIs, middlewares, Firebase, Supabase y Python."),
    ("Cloud & DevOps", "Google Cloud, Docker, Git, GitHub, GitLab, Bitbucket y flujos de entrega continua."),
    ("E-commerce & Data", "VTEX IO, checkout, facturación, Google Tag Manager, Meta Pixel, Google Ads y analítica."),
    ("Colaboración", "Scrum, Kanban, estimación técnica, consultoría y trabajo multidisciplinario con producto, diseño, marketing y negocio."),
]
for label, detail in skill_rows:
    p = style_paragraph(doc.add_paragraph(), after=3.5, line=1.1)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=9.55, color=DARK_PURPLE, bold=True)
    r = p.add_run(detail)
    set_run_font(r, size=9.55, color=INK)

add_section_heading(doc, "Educación y certificaciones", before=7)
p = style_paragraph(doc.add_paragraph(), after=2, line=1.05)
r = p.add_run("Administración en Tecnologías de la Información")
set_run_font(r, size=9.8, color=INK, bold=True)
r = p.add_run("  ·  Universidad Tecnológica de México · México")
set_run_font(r, size=9.5, color=MUTED)
p = style_paragraph(doc.add_paragraph(), after=2, line=1.05)
r = p.add_run("VTEX IO Developer")
set_run_font(r, size=9.8, color=INK, bold=True)
r = p.add_run("  ·  Certificación oficial VTEX")
set_run_font(r, size=9.5, color=MUTED)

add_section_heading(doc, "Propuesta de valor", before=7)
add_body(
    doc,
    "Combino especialización frontend con experiencia backend, cloud y e-commerce para desarrollar software completo: desde interfaces rápidas y mantenibles hasta APIs, automatizaciones e integraciones con servicios externos.",
    size=9.7,
    color=INK,
    after=0,
)

# Core document properties.
doc.core_properties.title = "CV — Uriel Mojica Mejía"
doc.core_properties.subject = "Senior Software Engineer · Frontend-first · Full-stack · E-commerce"
doc.core_properties.author = "Uriel Mojica Mejía"
doc.core_properties.keywords = "Software Engineer, React, Next.js, TypeScript, Node.js, Google Cloud, VTEX IO, E-commerce"
doc.core_properties.comments = "CV profesional actualizado"

doc.save(OUTPUT)
print(OUTPUT)
